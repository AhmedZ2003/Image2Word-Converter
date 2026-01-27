import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import pytesseract
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
import threading
import time

# If Tesseract is not in your PATH, uncomment and update:
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Set the theme
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class TechyOCRApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        # Window Setup
        self.title("Image2Word Converter")
        self.geometry("1100x700")
        
        # Grid Layout
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # Variables
        self.image_path = None
        self.current_doc_object = None  # Store the doc in memory before saving
        
        self.setup_sidebar()
        self.setup_main_area()

    def setup_sidebar(self):
        """Left panel with controls"""
        self.sidebar_frame = ctk.CTkFrame(self, width=200, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(5, weight=1) # Push bottom items down

        # Logo
        self.logo_label = ctk.CTkLabel(self.sidebar_frame, text="IMAGE 2\nWORD", 
                                      font=ctk.CTkFont(size=20, weight="bold"))
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 10))

        # Buttons
        self.btn_load = ctk.CTkButton(self.sidebar_frame, text="LOAD IMAGE", 
                                     command=self.select_image, 
                                     fg_color="#2b2b2b", hover_color="#3a3a3a", 
                                     border_width=2, border_color="#1f538d")
        self.btn_load.grid(row=1, column=0, padx=20, pady=10)

        self.btn_convert = ctk.CTkButton(self.sidebar_frame, text="INITIALIZE OCR", 
                                        command=self.start_conversion_thread, state="disabled",
                                        fg_color="#1f538d", hover_color="#14375e")
        self.btn_convert.grid(row=2, column=0, padx=20, pady=10)

        # SAVE BUTTON (Initially Disabled)
        self.btn_save = ctk.CTkButton(self.sidebar_frame, text="DOWNLOAD DOCX", 
                                     command=self.save_document, state="disabled",
                                     fg_color="#27ae60", hover_color="#2ecc71")
        self.btn_save.grid(row=3, column=0, padx=20, pady=10)

        # Progress Bar
        self.progress_bar = ctk.CTkProgressBar(self.sidebar_frame, width=150)
        self.progress_bar.grid(row=4, column=0, padx=20, pady=(10, 20))
        self.progress_bar.set(0)
        self.progress_bar.grid_remove()

        # Appearance Mode
        self.appearance_mode_label = ctk.CTkLabel(self.sidebar_frame, text="Interface Mode:", anchor="w")
        self.appearance_mode_label.grid(row=6, column=0, padx=20, pady=(10, 0))
        self.appearance_mode_optionemenu = ctk.CTkOptionMenu(self.sidebar_frame, values=["Dark", "Light"],
                                                                       command=self.change_appearance_mode_event)
        self.appearance_mode_optionemenu.grid(row=7, column=0, padx=20, pady=(10, 20))

    def setup_main_area(self):
        # Right panel with split view (Image | Text)
        self.main_frame = ctk.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.main_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        
        # Equal width columns
        self.main_frame.grid_columnconfigure((0, 1), weight=1, uniform="split")
        self.main_frame.grid_rowconfigure(0, weight=1)

        # Left: Image Preview ---
        self.preview_frame_l = ctk.CTkFrame(self.main_frame, corner_radius=10)
        self.preview_frame_l.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        
        self.lbl_img_title = ctk.CTkLabel(self.preview_frame_l, text="SOURCE INPUT", 
                                         font=ctk.CTkFont(size=12, weight="bold"), text_color="gray")
        self.lbl_img_title.pack(pady=10)
        
        self.lbl_image_preview = ctk.CTkLabel(self.preview_frame_l, text="No Data Loaded")
        self.lbl_image_preview.pack(expand=True)

        # Right: Text Result
        self.preview_frame_r = ctk.CTkFrame(self.main_frame, corner_radius=10)
        self.preview_frame_r.grid(row=0, column=1, sticky="nsew", padx=(10, 0))

        self.lbl_txt_title = ctk.CTkLabel(self.preview_frame_r, text="DIGITIZED OUTPUT", 
                                         font=ctk.CTkFont(size=12, weight="bold"), text_color="gray")
        self.lbl_txt_title.pack(pady=10)

        # Text Box
        self.textbox = tk.Text(self.preview_frame_r, font=("Consolas", 11), width=40,
                               bg="#2b2b2b", fg="#dce4ee", 
                               bd=0, highlightthickness=0, 
                               padx=15, pady=15, wrap="word")
        
        self.textbox.pack(fill="both", expand=True, padx=2, pady=(0, 10))
        self.textbox.insert("1.0", "System Idle.\nWaiting for input stream...")
        self.textbox.configure(state="disabled")

        # Status Footer
        self.status_label = ctk.CTkLabel(self.main_frame, text="STATUS: READY", anchor="w", text_color="#00ff00")
        self.status_label.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(10, 0))

    # LOGIC & EVENTS
    def change_appearance_mode_event(self, new_appearance_mode: str):
        ctk.set_appearance_mode(new_appearance_mode)

    def select_image(self):
        filename = filedialog.askopenfilename(filetypes=(("Images", "*.jpg;*.png;*.jpeg"),))
        if filename:
            self.image_path = filename
            self.display_image(filename)
            self.btn_convert.configure(state="normal")
            self.btn_save.configure(state="disabled") # Disable save when new image loads
            self.status_label.configure(text=f"STATUS: IMAGE LOADED", text_color="cyan")
            self.textbox.configure(state="normal")
            self.textbox.delete("1.0", "end")
            self.textbox.insert("1.0", ">> Image loaded.\n>> Press 'INITIALIZE OCR' to begin.")
            self.textbox.configure(state="disabled")

    def display_image(self, path):
        img = Image.open(path)
        # Smart Resize
        target_h = 400
        aspect = img.width / img.height
        new_w = int(target_h * aspect)
        
        img = img.resize((new_w, target_h), Image.Resampling.LANCZOS)
        self.ctk_image = ctk.CTkImage(light_image=img, dark_image=img, size=(new_w, target_h))
        self.lbl_image_preview.configure(image=self.ctk_image, text="")

    def start_conversion_thread(self):
        self.btn_convert.configure(state="disabled")
        self.btn_load.configure(state="disabled")
        self.btn_save.configure(state="disabled")
        self.progress_bar.grid()
        self.progress_bar.start()
        
        threading.Thread(target=self.run_ocr_process, daemon=True).start()

    def run_ocr_process(self):
        try:
            self.update_status("Scanning Geometry...", "yellow")
            time.sleep(0.5) 
            
            img = Image.open(self.image_path)
            hocr_data = pytesseract.image_to_pdf_or_hocr(img, extension='hocr').decode('utf-8')
            
            self.update_status("Parsing Formatting Tags...", "orange")
            words = self.parse_hocr(hocr_data)
            
            if not words:
                raise Exception("No readable text found.")

            self.update_status("Compiling Document...", "orange")
            # Store doc object in memory, don't save yet
            self.current_doc_object = self.generate_doc_object(words)
            
            self.after(0, lambda: self.conversion_complete())
            
        except Exception as e:
            self.after(0, lambda: self.conversion_failed(str(e)))

    def update_status(self, text, color):
        color_map = {"yellow": "#f1c40f", "orange": "#e67e22", "red": "#e74c3c", "green": "#2ecc71"}
        self.after(0, lambda: self.status_label.configure(text=f"STATUS: {text}", text_color=color_map.get(color, "white")))

    def conversion_complete(self):
        self.progress_bar.stop()
        self.progress_bar.grid_remove()
        self.btn_load.configure(state="normal")
        self.btn_convert.configure(state="normal")
        self.btn_save.configure(state="normal") # Enable Save Button
        
        self.update_status("COMPLETED. READY TO DOWNLOAD.", "green")
        
        # Display preview from memory
        self.display_text_result(self.current_doc_object)
        messagebox.showinfo("Success", "Conversion Complete!\n\nReview the preview on the right.\nClick 'DOWNLOAD DOCX' to save the file.")

    def conversion_failed(self, error_msg):
        self.progress_bar.stop()
        self.progress_bar.grid_remove()
        self.btn_load.configure(state="normal")
        self.btn_convert.configure(state="normal")
        self.update_status(f"ERROR: {error_msg}", "red")

    def save_document(self):
        if not self.current_doc_object:
            return
            
        # Open Save As Dialog
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"Converted_{os.path.basename(self.image_path).split('.')[0]}.docx"
        )
        
        if file_path:
            try:
                self.current_doc_object.save(file_path)
                messagebox.showinfo("Saved", f"File saved successfully at:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save file: {e}")

    # FORMATTING LOGIC
    def parse_hocr(self, hocr_string):
        words = []
        pattern = re.compile(r"<span class=['\"]ocrx_word['\"].*?title=['\"]bbox (\d+) (\d+) (\d+) (\d+).*?>(.*?)</span>", re.DOTALL)
        matches = pattern.findall(hocr_string)
        
        for match in matches:
            x1, y1, x2, y2, content = match
            x1, y1, x2, y2 = int(x1), int(y1), int(x2), int(y2)
            
            is_bold = bool(re.search(r"<strong>|<b>", content, re.IGNORECASE))
            is_italic = bool(re.search(r"<em>|<i>", content, re.IGNORECASE))
            clean_text = re.sub('<[^<]+?>', '', content).strip()
            
            if clean_text:
                words.append({'text': clean_text, 'x': x1, 'y': y1, 'h': y2-y1, 'bold': is_bold, 'italic': is_italic})
        return words

    def generate_doc_object(self, words_data):
        doc = Document()
        lines = {}
        for word in words_data:
            y = word['y']
            found = False
            for line_y in lines.keys():
                if abs(line_y - y) < 12: 
                    lines[line_y].append(word)
                    found = True
                    break
            if not found: lines[y] = [word]

        sorted_y = sorted(lines.keys())
        all_heights = [w['h'] for w in words_data]
        median_height = sorted(all_heights)[len(all_heights)//2] if all_heights else 20
        last_y_bottom = 0

        for y in sorted_y:
            line_words = sorted(lines[y], key=lambda k: k['x'])
            current_y_top = y
            gap = current_y_top - last_y_bottom
            
            if last_y_bottom > 0 and gap > (median_height * 1.5):
                doc.add_paragraph("") 

            p = doc.add_paragraph()
            if line_words[0]['x'] > 90: 
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            avg_h = sum([w['h'] for w in line_words]) / len(line_words)
            is_header = avg_h > (median_height * 1.3)

            for i, word in enumerate(line_words):
                text = word['text'] if i == 0 else " " + word['text']
                run = p.add_run(text)
                
                if is_header:
                    run.bold = True
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(11)
                    if word['bold']: run.bold = True
                    if word['italic']: run.italic = True
            
            max_h = max([w['h'] for w in line_words])
            last_y_bottom = y + max_h

        return doc # Return the object instead of saving path

    def display_text_result(self, doc_object):
        self.textbox.configure(state="normal")
        self.textbox.delete("1.0", "end")
        
        # Configure Tags
        self.textbox.tag_config("header", font=("Roboto", 14, "bold"), foreground="#62a1ff")
        self.textbox.tag_config("bold", font=("Roboto", 11, "bold"))
        self.textbox.tag_config("italic", font=("Roboto", 11, "italic"))
        self.textbox.tag_config("center", justify="center")

        for p in doc_object.paragraphs:
            tags = []
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                tags.append("center")
            
            if not p.text.strip():
                self.textbox.insert("end", "\n")
                continue

            for run in p.runs:
                run_tags = list(tags)
                if run.bold: run_tags.append("bold")
                if run.italic: run_tags.append("italic")
                if run.font.size and run.font.size.pt > 12: run_tags.append("header")
                
                self.textbox.insert("end", run.text, tuple(run_tags))
            
            self.textbox.insert("end", "\n", tuple(tags))
        
        self.textbox.configure(state="disabled")

if __name__ == "__main__":
    app = TechyOCRApp()
    app.mainloop()
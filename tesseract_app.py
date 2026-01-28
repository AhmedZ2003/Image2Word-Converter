import gradio as gr
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import tempfile

# CONFIGURATION: Set Tesseract path if needed
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def parse_hocr(hocr_string):
    words = []
    pattern = re.compile(r"<span class=['\"]ocrx_word['\"].*?title=['\"]bbox (\d+) (\d+) (\d+) (\d+).*?>(.*?)</span>", re.DOTALL)
    matches = pattern.findall(hocr_string)
    
    for match in matches:
        x1, y1, x2, y2, content = match
        x1, y1, x2, y2 = int(x1), int(y1), int(x2), int(y2)
        
        is_bold = bool(re.search(r"<strong>|<b>", content, re.IGNORECASE))
        is_italic = bool(re.search(r"<em>|<i>", content, re.IGNORECASE))
        clean_text = re.sub('<[^<]+?>', '', content).strip()
        
        if clean_text:
            words.append({'text': clean_text, 'x': x1, 'y': y1, 'h': y2-y1, 'bold': is_bold, 'italic': is_italic})
    return words

def generate_doc_and_preview(words_data):
    doc = Document()
    
    html_preview = "<div style='background-color: #ffffff; color: #ffffff; padding: 20px; font-family: monospace; border-radius: 5px;'>"
    
    lines = {}
    for word in words_data:
        y = word['y']
        found = False
        for line_y in lines.keys():
            if abs(line_y - y) < 12: 
                lines[line_y].append(word)
                found = True
                break
        if not found: lines[y] = [word]

    sorted_y = sorted(lines.keys())
    all_heights = [w['h'] for w in words_data]
    median_height = sorted(all_heights)[len(all_heights)//2] if all_heights else 20
    last_y_bottom = 0

    for y in sorted_y:
        line_words = sorted(lines[y], key=lambda k: k['x'])
        current_y_top = y
        gap = current_y_top - last_y_bottom
        
        # Spacing Logic
        if last_y_bottom > 0 and gap > (median_height * 1.5):
            doc.add_paragraph("")
            html_preview += "<br>"

        p = doc.add_paragraph()
        
        # Alignment Logic
        is_centered = line_words[0]['x'] > 90
        if is_centered: 
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            html_preview += "<div style='text-align: center;'>"
        else:
            html_preview += "<div>"

        avg_h = sum([w['h'] for w in line_words]) / len(line_words)
        is_header = avg_h > (median_height * 1.3)

        for i, word in enumerate(line_words):
            text = word['text'] if i == 0 else " " + word['text']
            run = p.add_run(text)
            
            # HTML Formatting for Preview
            html_word = text
            if is_header:
                run.bold = True
                run.font.size = Pt(14)
                # Headers remain light blue for distinction
                html_word = f"<span style='font-size: 1.3em; font-weight: bold; color: #62a1ff;'>{html_word}</span>"
            else:
                run.font.size = Pt(11)
                if word['bold']: 
                    run.bold = True
                    html_word = f"<b>{html_word}</b>"
                if word['italic']: 
                    run.italic = True
                    html_word = f"<i>{html_word}</i>"
            
            html_preview += html_word
            
        html_preview += "</div>"
        
        max_h = max([w['h'] for w in line_words])
        last_y_bottom = y + max_h
    
    html_preview += "</div>"
    return doc, html_preview

# GRADIO INTERFACE FUNCTION

def process_image(image):
    if image is None:
        return None, "<div style='color: red'>Please upload an image first.</div>"
    
    try:
        # 1. Tesseract OCR
        pil_img = Image.open(image)
        hocr_data = pytesseract.image_to_pdf_or_hocr(pil_img, extension='hocr').decode('utf-8')
        
        # 2. Parse
        words = parse_hocr(hocr_data)
        if not words:
            return None, "No text detected."
            
        # 3. Generate
        doc_obj, html_preview = generate_doc_and_preview(words)
        
        # 4. Save to temp file for download
        temp_dir = tempfile.gettempdir()
        filename = f"converted_doc_{os.urandom(4).hex()}.docx"
        save_path = os.path.join(temp_dir, filename)
        doc_obj.save(save_path)
        
        return save_path, html_preview

    except Exception as e:
        return None, f"Error: {str(e)}"

# UI LAYOUT
custom_css = """
body {background-color: #0b0f19;}
.gradio-container {font-family: 'Roboto', sans-serif;}
"""

with gr.Blocks(theme=gr.themes.Soft(primary_hue="blue", secondary_hue="slate"), css=custom_css, title="Image2Word") as app:
    
    gr.Markdown(
        """
        # 🧠 Image 2 Word Converter
        ### Professional Image-to-Word Converter with Formatting Preservation
        """
    )
    
    with gr.Row():
        with gr.Column(scale=1):
            img_input = gr.Image(type="filepath", label="Source Input", height=400)
            btn_run = gr.Button("INITIALIZE OCR", variant="primary")
        
        with gr.Column(scale=1):
            preview_output = gr.HTML(label="Digitized Preview", value="<div style='color:gray'>System Idle...</div>")
            file_output = gr.File(label="Download Result", interactive=False)

    btn_run.click(
        fn=process_image, 
        inputs=[img_input], 
        outputs=[file_output, preview_output]
    )

if __name__ == "__main__":
    app.launch(share=True)
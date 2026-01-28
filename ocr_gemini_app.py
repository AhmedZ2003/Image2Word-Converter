import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from PIL import Image
import os
import threading
import time
import base64
from io import BytesIO
import google.generativeai as genai

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Set the theme
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class TechyOCRApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        # Window Setup
        self.title("Image2Word OCR")
        self.geometry("1100x700")
        
        # Grid Layout
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # Variables
        self.image_path = None
        self.current_doc_object = None 

        self.api_key = os.getenv("OPENAI_API_KEY") 
        
        self.setup_sidebar()
        self.setup_main_area()

        # Check API Key on launch
        if not self.api_key:
            self.after(100, self.prompt_api_key)

    def prompt_api_key(self):
        key = simpledialog.askstring("OpenAI API Key", "Please enter your OpenAI API Key:", parent=self)
        if key:
            self.api_key = key
        else:
            messagebox.showwarning("Missing Key", "OCR features will not work without an API Key.")

    def setup_sidebar(self):
        """Left panel with controls"""
        self.sidebar_frame = ctk.CTkFrame(self, width=200, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(5, weight=1)

        # Logo
        self.logo_label = ctk.CTkLabel(self.sidebar_frame, text="Image 2\nWord", 
                                      font=ctk.CTkFont(size=20, weight="bold"))
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 10))

        # Buttons
        self.btn_load = ctk.CTkButton(self.sidebar_frame, text="LOAD IMAGE", 
                                     command=self.select_image, 
                                     fg_color="#2b2b2b", hover_color="#3a3a3a", 
                                     border_width=2, border_color="#1f538d")
        self.btn_load.grid(row=1, column=0, padx=20, pady=10)

        self.btn_convert = ctk.CTkButton(self.sidebar_frame, text="INITIALIZE AI", 
                                        command=self.start_conversion_thread, state="disabled",
                                        fg_color="#1f538d", hover_color="#14375e")
        self.btn_convert.grid(row=2, column=0, padx=20, pady=10)

        self.btn_save = ctk.CTkButton(self.sidebar_frame, text="DOWNLOAD DOCX", 
                                     command=self.save_document, state="disabled",
                                     fg_color="#27ae60", hover_color="#2ecc71")
        self.btn_save.grid(row=3, column=0, padx=20, pady=10)

        # Progress Bar
        self.progress_bar = ctk.CTkProgressBar(self.sidebar_frame, width=150)
        self.progress_bar.grid(row=4, column=0, padx=20, pady=(10, 20))
        self.progress_bar.set(0)
        self.progress_bar.grid_remove()

        # Appearance Mode
        self.appearance_mode_label = ctk.CTkLabel(self.sidebar_frame, text="Interface Mode:", anchor="w")
        self.appearance_mode_label.grid(row=6, column=0, padx=20, pady=(10, 0))
        self.appearance_mode_optionemenu = ctk.CTkOptionMenu(self.sidebar_frame, values=["Dark", "Light"],
                                                               command=self.change_appearance_mode_event)
        self.appearance_mode_optionemenu.grid(row=7, column=0, padx=20, pady=(10, 20))

    def setup_main_area(self):
        self.main_frame = ctk.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.main_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        
        self.main_frame.grid_columnconfigure((0, 1), weight=1, uniform="split")
        self.main_frame.grid_rowconfigure(0, weight=1)

        # Left: Image Preview
        self.preview_frame_l = ctk.CTkFrame(self.main_frame, corner_radius=10)
        self.preview_frame_l.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        
        self.lbl_img_title = ctk.CTkLabel(self.preview_frame_l, text="SOURCE INPUT", 
                                        font=ctk.CTkFont(size=12, weight="bold"), text_color="gray")
        self.lbl_img_title.pack(pady=10)
        
        self.lbl_image_preview = ctk.CTkLabel(self.preview_frame_l, text="No Data Loaded")
        self.lbl_image_preview.pack(expand=True)

        # Right: Text Result
        self.preview_frame_r = ctk.CTkFrame(self.main_frame, corner_radius=10)
        self.preview_frame_r.grid(row=0, column=1, sticky="nsew", padx=(10, 0))

        self.lbl_txt_title = ctk.CTkLabel(self.preview_frame_r, text="OUTPUT", 
                                        font=ctk.CTkFont(size=12, weight="bold"), text_color="gray")
        self.lbl_txt_title.pack(pady=10)

        # Text Box
        self.textbox = tk.Text(self.preview_frame_r, font=("Consolas", 11), width=40,
                               bg="#2b2b2b", fg="#dce4ee", 
                               bd=0, highlightthickness=0, 
                               padx=15, pady=15, wrap="word")
        
        self.textbox.pack(fill="both", expand=True, padx=2, pady=(0, 10))
        self.textbox.insert("1.0", "System Idle.\nWaiting for input stream...")
        self.textbox.configure(state="disabled")

        # Status Footer
        self.status_label = ctk.CTkLabel(self.main_frame, text="STATUS: READY", anchor="w", text_color="#00ff00")
        self.status_label.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(10, 0))

    # LOGIC & EVENTS
    def change_appearance_mode_event(self, new_appearance_mode: str):
        ctk.set_appearance_mode(new_appearance_mode)

    def select_image(self):
        filename = filedialog.askopenfilename(filetypes=(("Images", "*.jpg;*.png;*.jpeg;*.webp"),))
        if filename:
            self.image_path = filename
            self.display_image(filename)
            self.btn_convert.configure(state="normal")
            self.btn_save.configure(state="disabled")
            self.status_label.configure(text=f"STATUS: IMAGE LOADED", text_color="cyan")
            self.textbox.configure(state="normal")
            self.textbox.delete("1.0", "end")
            self.textbox.insert("1.0", ">> Image loaded.\n>> Press 'INITIALIZE AI' to send to Gemini 2.5")
            self.textbox.configure(state="disabled")

    def display_image(self, path):
        img = Image.open(path)
        target_h = 400
        aspect = img.width / img.height
        new_w = int(target_h * aspect)
        
        img = img.resize((new_w, target_h), Image.Resampling.LANCZOS)
        self.ctk_image = ctk.CTkImage(light_image=img, dark_image=img, size=(new_w, target_h))
        self.lbl_image_preview.configure(image=self.ctk_image, text="")

    def start_conversion_thread(self):
        if not self.api_key:
            messagebox.showerror("Error", "No API Key found.")
            self.prompt_api_key()
            return

        self.btn_convert.configure(state="disabled")
        self.btn_load.configure(state="disabled")
        self.btn_save.configure(state="disabled")
        self.progress_bar.grid()
        self.progress_bar.start()
        
        threading.Thread(target=self.run_ocr_process, daemon=True).start()

    # --- REPLACED OCR LOGIC WITH GPT LOGIC ---
    def encode_image(self, image_path):
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def run_ocr_process(self):
        try:
            self.update_status("Configuring Gemini AI...", "yellow")
            
            # Configure the Google API
            genai.configure(api_key=self.api_key)
            
            # Load the Model
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            self.update_status("Processing Image...", "orange")
            
            # Load image directly with PIL (Google handles PIL images natively)
            img = Image.open(self.image_path)
            
            # The Prompt
            prompt = "Extract the text from this image. Return the content in Markdown format. Use headers (#) for big text, bold (**) for bold text. Do not include markdown code block fences. Just return the raw text."
            
            # Send to Gemini
            response = model.generate_content([prompt, img])
            
            # Get text
            result_text = response.text
            
            self.update_status("Compiling Document...", "orange")
            
            # Generate Doc Object from the Markdown String
            self.current_doc_object = self.markdown_to_docx(result_text)
            
            # Pass the raw text to completion for display
            self.after(0, lambda text=result_text: self.conversion_complete(text))
            
        except Exception as e:
            # Convert 'e' to a string immediately
            error_msg = str(e)
            self.after(0, lambda: self.conversion_failed(error_msg))

    def update_status(self, text, color):
        color_map = {"yellow": "#f1c40f", "orange": "#e67e22", "red": "#e74c3c", "green": "#2ecc71"}
        self.after(0, lambda: self.status_label.configure(text=f"STATUS: {text}", text_color=color_map.get(color, "white")))

    def conversion_complete(self, raw_text):
        self.progress_bar.stop()
        self.progress_bar.grid_remove()
        self.btn_load.configure(state="normal")
        self.btn_convert.configure(state="normal")
        self.btn_save.configure(state="normal")
        
        self.update_status("COMPLETED. READY TO DOWNLOAD.", "green")
        
        # Display preview
        self.display_text_result(raw_text)
        messagebox.showinfo("Success", "AI Conversion Complete!")

    def conversion_failed(self, error_msg):
        self.progress_bar.stop()
        self.progress_bar.grid_remove()
        self.btn_load.configure(state="normal")
        self.btn_convert.configure(state="normal")
        self.update_status(f"ERROR: {error_msg}", "red")

    def save_document(self):
        if not self.current_doc_object:
            return
            
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"GPT_Converted_{os.path.basename(self.image_path).split('.')[0]}.docx"
        )
        
        if file_path:
            try:
                self.current_doc_object.save(file_path)
                messagebox.showinfo("Saved", f"File saved successfully at:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save file: {e}")

    # FORMATTING LOGIC (Markdown -> Docx)
    def markdown_to_docx(self, text):
        doc = Document()
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Header Handling (# Header)
            if line.startswith('#'):
                level = line.count('#')
                clean_text = line.replace('#', '').strip()
                p = doc.add_heading(clean_text, level=min(level, 3) if level < 9 else 1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # Standard Paragraph with simple Bold parsing
                p = doc.add_paragraph()
                
                # Very basic markdown bold parser (**text**)
                # Splitting by ** gives: ["Regular ", "Bold", " Regular"]
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
                
                p.style = doc.styles['Normal']
                
        return doc

    def display_text_result(self, raw_text):
        self.textbox.configure(state="normal")
        self.textbox.delete("1.0", "end")
        
        # Configure Tags
        self.textbox.tag_config("header", font=("Roboto", 14, "bold"), foreground="#62a1ff")
        
        lines = raw_text.split('\n')
        for line in lines:
            if line.startswith('#'):
                self.textbox.insert("end", line + "\n", "header")
            else:
                self.textbox.insert("end", line + "\n")
        
        self.textbox.configure(state="disabled")

if __name__ == "__main__":
    app = TechyOCRApp()
    app.mainloop()
import gradio as gr
import google.generativeai as genai
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import tempfile

def markdown_to_docx(text):
    # Converts Markdown text (headers and bold) into a DOCX object.
    doc = Document()
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Header Handling (# Header)
        if line.startswith('#'):
            level = line.count('#')
            clean_text = line.replace('#', '').strip()
            # Docx supports heading levels 1-9
            heading_level = min(level, 9)
            p = doc.add_heading(clean_text, level=heading_level)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Standard Paragraph with simple Bold parsing
            p = doc.add_paragraph()
            
            # Very basic markdown bold parser (**text**)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            
            p.style = doc.styles['Normal']
            
    return doc

def process_image(image, api_key):
    # Takes an image and API key, returns the raw text and a path to the .docx file.
    if image is None:
        return "Please upload an image.", None
    
    if not api_key:
        return "Please enter a valid Google Gemini API Key.", None

    try:
        # 1. Configure API
        genai.configure(api_key=api_key)
        
        # 2. Select Model
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        # 3. Define Prompt
        prompt = (
            "Extract the text from this image. Return the content in Markdown format. "
            "Use headers (#) for big text, bold (**) for bold text. "
            "Do not include markdown code block fences (like ```markdown). "
            "Just return the raw text."
        )
        
        
        # 4. Call Gemini
        response = model.generate_content([prompt, image])
        result_text = response.text
        
        # 5. Generate DOCX
        doc = markdown_to_docx(result_text)
        
        # 6. Save to a temporary file
        # Hugging Face spaces act like read-only containers mostly, 
        # so we use a temporary file path for the output.
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        temp_file.close()
        
        return result_text, temp_file.name

    except Exception as e:
        return f"Error: {str(e)}", None

# Gradio Interface Setup

# Custom CSS to make it look a bit cleaner
custom_css = """
#component-0 {max_width: 800px; margin: auto;}
"""

with gr.Blocks(css=custom_css, title="AI Image to Word Converter") as demo:
    gr.Markdown("# 📄 AI Image to Word (Docx) Converter")
    gr.Markdown("Upload an image, enter your Google Gemini API Key, and get a formatted Word document back.")
    
    with gr.Row():
        with gr.Column():
            api_input = gr.Textbox(
                label="Google Gemini API Key", 
                type="password", 
                placeholder="Paste your key here (starts with AIza...)"
            )
            image_input = gr.Image(type="pil", label="Upload Image")
            submit_btn = gr.Button("🚀 Convert Image", variant="primary")
        
        with gr.Column():
            output_text = gr.TextArea(label="Extracted Text (Markdown)", interactive=False)
            output_file = gr.File(label="Download DOCX")

    # Connect the button to the function
    submit_btn.click(
        fn=process_image, 
        inputs=[image_input, api_input], 
        outputs=[output_text, output_file]
    )
    
    gr.Markdown("Powered by **Gemini 2.5 Flash**")

# Launch the app
if __name__ == "__main__":
    demo.launch()